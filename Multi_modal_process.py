import os
import base64
from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader
import PyPDF2
from docx import Document
import tempfile

def image_process(image):
    image_base64 = base64.b64encode(image.getvalue()).decode("utf-8")
    return f"data:{image.type};base64,{image_base64}"


def show_uploaded_file(user_file):
    text = ""
    if user_file.name.endswith((".txt", ".md")):
        text = user_file.read().decode("utf-8")
    elif user_file.name.endswith(".pdf"):
        pdf_reader = PyPDF2.PdfReader(user_file)
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
    elif user_file.name.endswith(".docx"):
        doc = Document(user_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
    return text.strip()


def ui_file_process(user_files):
    a = []
    for i in user_files:
        text = show_uploaded_file(i)
        a.append({"filename": i.name, "text": text})
    return a

def file_stream_to_path(uploaded_file):
    # 获取文件后缀，例如 .pdf / .txt / .docx
    file_suffix = os.path.splitext(uploaded_file.name)[1]
    # 创建临时文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_suffix) as tmp_file:
        # 把上传文件的二进制内容写入临时文件
        tmp_file.write(uploaded_file.getvalue())
        # 返回临时文件路径
        return tmp_file.name


def load_documents(file_path, filename):
    filename = filename.lower()
    if filename.endswith(".pdf"):
        loader = PyPDFLoader(file_path)
    elif filename.endswith(".txt") or filename.endswith(".md"):
        loader = TextLoader(file_path, encoding="utf-8")
    elif filename.endswith(".docx"):
        loader = Docx2txtLoader(file_path)
    else:
        raise ValueError(f"暂不支持的文件类型：{filename}")
    docs = loader.load()
    return docs